#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Aug 24 13:35:53 2023

@author: kieranmartin
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import datetime
import os
from io import BytesIO
import requests
import re
from collections import Counter
import qualtrics_api_toolkit.qual_api as qa
from qualtrics_api_toolkit.utils import (
    QUALTRICS_CREDS,
    get_token,
)
    
# Extract client ID, secret, and data center from credentials
client_id = QUALTRICS_CREDS.get('ID')
client_secret = QUALTRICS_CREDS.get('Secret')
data_center = QUALTRICS_CREDS.get('DataCenter')
base_url = f'https://{data_center}.qualtrics.com'

grant_type = 'client_credentials'
scope = 'read:surveys read:survey_responses'
data = qa.return_kwargs_as_dict(grant_type=grant_type, scope=scope)


# Get the bearer token
token = get_token(base_url, client_id, client_secret, data)

# Retrieve the list of surveys and find the survey ID
survey_list_df = qa.get_survey_list(base_url, token)
# select survey from tkinter generated dropdown list of accounts surveys
survey_name = qa.select_survey(survey_list_df)
survey_id = qa.get_survey_id_by_name(survey_list_df, survey_name)

# Export survey responses and track the progress
response_export = qa.start_response_export(base_url, token, survey_id)
export_progress_id = response_export.get('result').get('progressId')

# Wait for export to complete and retrieve file ID for responses
file_id = qa.wait_for_export_completion(base_url, token, survey_id, export_progress_id)

# Download and process the survey responses
survey_responses = qa.get_response_export_file(base_url, token, survey_id, file_id)

responses_df = qa.extract_and_organize_responses(survey_responses)
responses_df = qa.filter_preview_responses(responses_df)

# Retrieve survey questions and map them to response columns
survey = qa.get_survey_questions(base_url, token, survey_id).get('result')
survey_questions = survey.get('questions')

blocks_df = qa.get_block_data(survey)

question_df, question_values_df = qa.extract_column_data_types(
    survey_questions, 
    responses_df, 
    base_url, 
    token, 
    survey_id
)
# Remove timer questions after creating question_df. could edit handle_timing_question
# to remove timer questions before they are organized into question_df
question_df = question_df[question_df['question_selector'] != 'PageTimer']
question_df = qa.create_data_type_dictionary(question_df, question_values_df)
question_df = qa.reorder_question_df_with_normalized_ids(question_df, blocks_df, drop_na=True)

# Clean responses and retain only question columns
responses_df = qa.clean_responses(responses_df, question_df)

# Filter question_df for FreeText columns and exclude those containing 'TEXT_'
free_text_columns = question_df[question_df.data_type == 'FreeText']

# Map question_id to question_text
question_id_to_text = free_text_columns.set_index('question_id')['question_text']

# Create text_df using question_id, then rename columns to question_text
text_df = responses_df[question_id_to_text.index]
text_df.columns = text_df.columns.map(question_id_to_text)


def process_response_column(responses_df, column):
    """
    Cleans and processes a multiple-choice or rank order response column by checking types 
    and handling lists, NaNs, and other cases as in the original structure.
    
    Parameters:
    responses_df (pd.DataFrame): Survey responses DataFrame.
    column (str): Column name for the multiple-choice or rank order question.
    
    Returns:
    list: Processed list of responses with 'NULL' for NaNs.
    """
    new_list = []
    
    # Check if the column is of object type
    if responses_df[column].dtype == 'object':
        for response in responses_df[column]:
            if isinstance(response, list):
                # Add list responses directly
                new_list += response
            elif isinstance(response, float) and np.isnan(response):
                new_list.append('NULL')
            elif isinstance(response, str) and response.isdigit():
                new_list.append(response)  # Keep numeric string values
            else:
                # Log unexpected values for detailed debugging
                print(f"Unexpected value in {column}: {response} (Type: {type(response)})")
                new_list.append('NULL')
    
    elif responses_df[column].dtype in ['float', 'int64', 'int']:  # Catch int64 as well
        for response in responses_df[column]:
            if isinstance(response, float) and np.isnan(response):
                new_list.append('NULL')
            else:
                new_list.append(str(int(response)))  # Convert to int and then to string
    else:
        print(f"Problems with conversion in column {column} (Unexpected data type: {responses_df[column].dtype})")
    
    return new_list



def process_multiple_choice_question(responses_df, col, question_values_df):
    """
    Processes a multiple-choice question to generate a frequency table.
    """
    new_list = process_response_column(responses_df, col)
    new_list = [response for response in new_list if response != 'NULL']

    freq_dist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
    freq_df.columns = ['answer_id', 'N']
    
    current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
    temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
    temp_df['N'] = temp_df['N'].astype(int)
    temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
    temp_df.columns = ['Code', 'Value', 'Count', 'Frequency']
    
    # Add line breaks to 'Value' for better x-tick display
    temp_df['Value'] = temp_df['Value'].apply(
        lambda x: (
            ' '.join(x.split()[:3]) + '\n' + ' '.join(x.split()[3:])
            if len(x.split()) > 4 else
            (x.split()[0] + '\n' + ' '.join(x.split()[1:])
             if len(x.split()) < 4 and len(x) > 13 else x)
        )
    )
    
    return temp_df


def process_numeric_question(responses_df, col, question_values_df):
    """
    Processes a numeric question to generate a frequency table.
    """
    new_list = process_response_column(responses_df, col)
    
    # Exclude 'NULL'
    new_list = [response for response in new_list if response != 'NULL']

    freq_dist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
    freq_df.columns = ['answer_id', 'N']
    
    current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
    temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
    temp_df['N'] = temp_df['N'].astype(int)
    temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
    temp_df.columns = ['Value', 'Value_Label', 'Count', 'Frequency']
    
    # Drop the 'Value_Label' column
    temp_df = temp_df[['Value', 'Count', 'Frequency']]
    
    # Add line breaks to 'Value' for better x-tick display
    temp_df['Value'] = temp_df['Value'].apply(
        lambda x: (
            ' '.join(x.split()[:3]) + '\n' + ' '.join(x.split()[3:])
            if len(x.split()) > 4 else
            (x.split()[0] + '\n' + ' '.join(x.split()[1:])
             if len(x.split()) < 4 and len(x) > 13 else x)
        )
    )
    
    return temp_df


def add_chart_to_doc(doc, result_df, is_numeric=False):
    """
    Adds a bar chart to the document for each question.
    The y-axis represents counts, and bar labels show the count of responses.
    The y-axis ticks are set to [5, 10, 15, 20, 25, 30].
    """
    chart_label = 'Value'

    # Create a bar chart if the result has fewer than 10 entries
    if len(result_df) < 10:
        fig, ax = plt.subplots(figsize=(7, 3))

        # Use count values for the y-axis
        count_values = result_df['Count']
        bar_colors = ['dodgerblue' if val != 'NULL' else 'crimson' for val in result_df['Value']]
        ax.bar(result_df['Value'], count_values, color=bar_colors, width=0.4, zorder=3)

        # Set x-axis and y-axis labels
        #ax.set_xlabel(chart_label, labelpad=10, ha='center')  # Center-align x-axis label
        #ax.set_ylabel('Count', labelpad=10, ha='center')  # Center-align y-axis label
        ax.set_xticks(range(len(result_df['Value'])))
        ax.set_yticks([5, 10, 15, 20])
        
        # Add light gray dotted horizontal gridlines
        ax.grid(axis='y', color='lightgray', linestyle='dotted', linewidth=0.75, zorder=1)

        # Annotate bars with count labels (centered above the bars)
        for p, count in zip(ax.patches, count_values):
            ax.annotate(
                str(count),  # The label text, which is the count value
                (p.get_x() + p.get_width() / 2., p.get_height()),  # The position (x,y)
                ha='center',  # Center alignment of text
                va='bottom',  # Align the bottom of the text at the given position
                xytext=(0, 10),  # 10 points vertical offset
                textcoords='offset points',  # Offset (in points) from the specified position
                fontsize=10,  # Font size of the text
            )

        # Adjust layout to make plot content larger
        plt.subplots_adjust(left=0.2, right=0.95, top=0.9, bottom=0.2)  # Adjust plot area margins
        plt.tight_layout()  # Automatically optimize spacing

        # Save the chart as an image and add to the document
        image_stream = BytesIO()
        plt.savefig(image_stream, format='png')
        plt.close(fig)
        image_stream.seek(0)
        doc.add_picture(image_stream, width=Inches(7), height=Inches(3))
    
    #doc.add_page_break()

    return doc


def generate_response_frequency(responses_df, question_df, question_values_df):
    """
    Generates frequency plots for multiple-choice and numeric-choice questions,
    adds FreeText responses, and groups questions by block names.
    """
    result_dict = {}
    numeric_result_dict = {}

    multiple_choice_columns = list(question_df.question_id[question_df["data_type"] == 'MultipleChoice'])
    for col in multiple_choice_columns:
        result_dict[col] = process_multiple_choice_question(responses_df, col, question_values_df)
    
    numeric_choice_columns = list(question_df.question_id[question_df["data_type"] == 'Numeric'])
    for col in numeric_choice_columns:
        numeric_result_dict[col] = process_numeric_question(responses_df, col, question_values_df)
    
    # Filter question_df to include only processed questions
    processed_questions = set(result_dict.keys()).union(set(numeric_result_dict.keys()))
    filtered_question_df = question_df[question_df['question_id'].isin(processed_questions)]
    
    # Document setup
    N = len(responses_df)
    number_of_responses_text = f"N = {N} responses"
    doc = Document()
    
    # Set custom margins
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Add survey name to the header
    header = section.header
    header_paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_paragraph.text = survey_name  # Replace with the actual survey name
    header_paragraph.style = 'Normal'
    #header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Replace with actual survey name
    doc.add_paragraph(number_of_responses_text, style='Normal')
    
    # Group questions by block
    current_block = None
    for _, row in filtered_question_df.iterrows():
        block_name = row['Block Name']
        if block_name != current_block:
            # Add block name as a heading
            doc.add_heading(block_name, level=1)
            current_block = block_name
        
        # Add question details
        col = row['question_id']
        question_name = row['question_name']
        question_text = row['question_text']
        mapped_text = f"{question_name}: {question_text}"
        
        doc.add_heading(mapped_text, level=2)
        
        # Add the total number of responses as a paragraph below the question text
        total_responses = len(responses_df[col].dropna())  # Drop NaNs to count actual responses
        response_text = f"{total_responses} Responses"
        doc.add_paragraph(response_text, style='Normal')
        
        if row["data_type"] == 'MultipleChoice' and col in result_dict:
            add_chart_to_doc(doc, result_dict[col])
        elif row["data_type"] == 'Numeric' and col in numeric_result_dict:
            add_chart_to_doc(doc, numeric_result_dict[col], is_numeric=True)
    
    # Add FreeText responses
    free_text_columns = question_df[question_df.data_type == 'FreeText']
    question_id_to_name_text = free_text_columns.set_index('question_id')[['question_name', 'question_text', 'Block Name']].to_dict('index')

    for question_id, name_text in question_id_to_name_text.items():
        block_name = name_text['Block Name']
        if block_name != current_block:
            # Add block name as a heading
            doc.add_heading(block_name, level=1)
            current_block = block_name
        
        question_name, question_text = name_text['question_name'], name_text['question_text']
        
        doc.add_heading(f"{question_name}: {question_text}", level=2)
        
        responses = responses_df[question_id].dropna().tolist()
        
        if responses:
            for response in responses:
                doc.add_paragraph(f"- {response}", style='List Bullet')
        else:
            doc.add_paragraph("No responses.", style='Normal')
    
    return doc


doc = generate_response_frequency(responses_df, question_df, question_values_df)
doc.save('C:\\Users\\484843\\Documents\\GitHub\\test.docx')